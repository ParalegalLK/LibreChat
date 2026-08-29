#!/usr/bin/env python3
"""
Post-processor: convert plain-text paragraph numbers into Word auto-numbering.

Scans a Pandoc-generated DOCX for paragraphs whose text begins with:
  - "N. "   → level 0 (decimal numbering: 1. 2. 3.)
  - "(x) "  → level 1 (lowercase alpha in parens: (a) (b) (c))

For each match it:
  1. Strips the text prefix from the first run (preserving formatting)
  2. Applies <w:numPr> so Word treats it as an auto-numbered paragraph

Only processes "First Paragraph" and "Body Text" styled paragraphs.
"""

import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ELIGIBLE_STYLES = {"First Paragraph", "Body Text"}
HEADING_STYLES = {"Heading 1": 0, "Heading 2": 1}

# Regex for the text prefix at the start of a paragraph
RE_LEVEL0 = re.compile(r"^(\d+)\.\s")  # "42. "
RE_LEVEL1 = re.compile(r"^\([a-z]+\)\s")  # "(a) "
RE_HEADING_NUM = re.compile(r"^[\d.]+\s")  # "1 " or "2.1 "


def create_numbering_definition(numbering_part):
    """Add a multi-level abstract numbering definition and a concrete num."""
    numbering_elem = numbering_part._element

    # Pick IDs that won't clash with existing definitions
    existing_abstract_ids = {
        int(a.get(qn("w:abstractNumId")))
        for a in numbering_elem.findall(qn("w:abstractNum"))
    }
    abstract_id = max(existing_abstract_ids, default=-1) + 1

    existing_num_ids = {
        int(n.get(qn("w:numId")))
        for n in numbering_elem.findall(qn("w:num"))
    }
    num_id = max(existing_num_ids, default=0) + 1

    # --- abstractNum ---
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level)

    # Level 0: decimal  "1."
    lvl0 = _make_level(
        ilvl=0,
        num_fmt="decimal",
        lvl_text="%1.",
        left=720,
        hanging=720,
        start=1,
    )
    abstract_num.append(lvl0)

    # Level 1: lowerLetter "(a)"
    lvl1 = _make_level(
        ilvl=1,
        num_fmt="lowerLetter",
        lvl_text="(%2)",
        left=1440,
        hanging=720,
        start=1,
        lvl_restart=1,  # restart at (a) when level 0 increments
    )
    abstract_num.append(lvl1)

    # Insert abstractNum before any <w:num> elements (schema requirement)
    first_num = numbering_elem.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elem.append(abstract_num)

    # --- num ---
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id_elem = OxmlElement("w:abstractNumId")
    abstract_num_id_elem.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id_elem)
    numbering_elem.append(num)

    return num_id


def create_heading_numbering_definition(numbering_part):
    """Add a multi-level numbering definition for section headings."""
    numbering_elem = numbering_part._element

    existing_abstract_ids = {
        int(a.get(qn("w:abstractNumId")))
        for a in numbering_elem.findall(qn("w:abstractNum"))
    }
    abstract_id = max(existing_abstract_ids, default=-1) + 1

    existing_num_ids = {
        int(n.get(qn("w:numId")))
        for n in numbering_elem.findall(qn("w:num"))
    }
    num_id = max(existing_num_ids, default=0) + 1

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

    multi_level = OxmlElement("w:multiLevelType")
    multi_level.set(qn("w:val"), "multilevel")
    abstract_num.append(multi_level)

    # Level 0: Heading 1 → "1"
    lvl0 = _make_level(
        ilvl=0, num_fmt="decimal", lvl_text="%1",
        left=0, hanging=0, start=1, suff="space",
    )
    abstract_num.append(lvl0)

    # Level 1: Heading 2 → "1.1"
    lvl1 = _make_level(
        ilvl=1, num_fmt="decimal", lvl_text="%1.%2",
        left=0, hanging=0, start=1, lvl_restart=1, suff="space",
    )
    abstract_num.append(lvl1)

    first_num = numbering_elem.find(qn("w:num"))
    if first_num is not None:
        first_num.addprevious(abstract_num)
    else:
        numbering_elem.append(abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id_elem = OxmlElement("w:abstractNumId")
    abstract_num_id_elem.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id_elem)
    numbering_elem.append(num)

    return num_id


def _make_level(ilvl, num_fmt, lvl_text, left, hanging, start=1, lvl_restart=None, suff=None):
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), str(ilvl))

    start_elem = OxmlElement("w:start")
    start_elem.set(qn("w:val"), str(start))
    lvl.append(start_elem)

    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), num_fmt)
    lvl.append(fmt)

    if suff is not None:
        suff_elem = OxmlElement("w:suff")
        suff_elem.set(qn("w:val"), suff)
        lvl.append(suff_elem)

    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), lvl_text)
    lvl.append(text)

    jc = OxmlElement("w:lvlJc")
    jc.set(qn("w:val"), "left")
    lvl.append(jc)

    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    ppr.append(ind)
    lvl.append(ppr)

    if lvl_restart is not None:
        restart = OxmlElement("w:lvlRestart")
        restart.set(qn("w:val"), str(lvl_restart))
        # Insert after w:start
        lvl.insert(1, restart)

    return lvl


def apply_num_pr(paragraph, num_id, ilvl):
    """Set <w:numPr> on a paragraph's pPr."""
    pPr = paragraph._element.get_or_add_pPr()

    # Remove any existing numPr
    old = pPr.find(qn("w:numPr"))
    if old is not None:
        pPr.remove(old)

    numPr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), str(ilvl))
    numPr.append(ilvl_elem)
    numId_elem = OxmlElement("w:numId")
    numId_elem.set(qn("w:val"), str(num_id))
    numPr.append(numId_elem)

    # Insert numPr at beginning of pPr (before other properties)
    pPr.insert(0, numPr)


def strip_text_prefix(paragraph, pattern):
    """Remove the matching prefix from paragraph runs.

    The prefix may span multiple runs (e.g. run0="65." run1=" " run2="text").
    We concatenate leading runs until we find the full pattern match, then
    strip the matched portion across those runs.
    """
    runs = paragraph.runs
    if not runs:
        return False

    # Try matching against progressively more runs
    concat = ""
    for i, run in enumerate(runs):
        concat += run.text
        m = pattern.match(concat)
        if m:
            # The match spans runs 0..i — strip the matched text
            remaining = concat[m.end():]
            # Delete all fully-consumed runs (0..i-1)
            for j in range(i):
                runs[j].text = ""
            # Set the last matched run to whatever remains after the match
            runs[i].text = remaining
            # Clean up empty leading runs and stray spaces
            _clean_leading_runs(runs)
            return True
        # If we already have enough text that the pattern can't possibly match
        # at the start, bail early
        if len(concat) > 20:
            break

    return False


def _clean_leading_runs(runs):
    """Remove empty leading runs and collapse double-spaces."""
    # Find first non-empty run
    for i, run in enumerate(runs):
        if run.text == "":
            continue
        # If the first non-empty run starts with space and the previous run
        # was emptied (i.e. i > 0), strip that leading space
        if i > 0 and run.text.startswith(" "):
            run.text = run.text[1:]
        break


def process_document(docx_path):
    doc = Document(docx_path)

    # Access numbering part (Pandoc usually creates one; if not, create it)
    numbering_part = doc.part.numbering_part
    num_id = create_numbering_definition(numbering_part)

    level0_count = 0
    level1_count = 0

    for para in doc.paragraphs:
        if para.style.name not in ELIGIBLE_STYLES:
            continue

        text = para.text

        # Check level 0 first (more specific: starts with digit)
        if RE_LEVEL0.match(text):
            if strip_text_prefix(para, RE_LEVEL0):
                apply_num_pr(para, num_id, ilvl=0)
                level0_count += 1

        elif RE_LEVEL1.match(text):
            if strip_text_prefix(para, RE_LEVEL1):
                apply_num_pr(para, num_id, ilvl=1)
                level1_count += 1

    # Process headings
    heading_num_id = create_heading_numbering_definition(numbering_part)
    heading_count = 0

    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name not in HEADING_STYLES:
            continue
        ilvl = HEADING_STYLES[style_name]
        if strip_text_prefix(para, RE_HEADING_NUM):
            apply_num_pr(para, heading_num_id, ilvl=ilvl)
            heading_count += 1

    doc.save(docx_path)
    print(f"Applied auto-numbering: {level0_count} paragraphs, {level1_count} sub-paragraphs, {heading_count} headings")


def main():
    if len(sys.argv) != 2:
        print(f"Usage: {sys.argv[0]} <docx_file>", file=sys.stderr)
        sys.exit(1)

    docx_path = Path(sys.argv[1])
    if not docx_path.exists():
        print(f"Error: {docx_path} not found", file=sys.stderr)
        sys.exit(1)

    process_document(str(docx_path))


if __name__ == "__main__":
    main()
